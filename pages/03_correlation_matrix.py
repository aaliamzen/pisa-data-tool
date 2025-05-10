import streamlit as st
import pandas as pd
import numpy as np
import re
import streamlit.components.v1 as components
import matplotlib.pyplot as plt
from scipy.stats import t
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import textwrap
import itertools

# Add logo that persists across all pages
try:
    st.logo("assets/logo.png")  # Replace with the path to your logo file, e.g., "assets/logo.png"
except Exception as e:
    st.error(f"Failed to load logo: {e}")

# Streamlit app configuration
st.set_page_config(page_title="Correlation Matrix - PISA Data Exploration Tool", layout="wide")

# Function to compute weighted correlation coefficient
def weighted_correlation(x, y, w):
    try:
        if len(x) != len(y) or len(x) != len(w):
            raise ValueError("Input arrays must have the same length.")
        if not all(w >= 0):
            raise ValueError("Weights must be non-negative.")
        if len(x) <= 1:
            return np.nan, np.nan
        # Remove NaN values
        mask = ~(np.isnan(x) | np.isnan(y) | np.isnan(w))
        x = x[mask]
        y = y[mask]
        w = w[mask]
        if len(x) <= 1:
            return np.nan, np.nan
        # Weighted mean
        w_sum = np.sum(w)
        if w_sum == 0:
            return np.nan, np.nan
        wx_mean = np.sum(w * x) / w_sum
        wy_mean = np.sum(w * y) / w_sum
        # Weighted covariance and variances
        cov_xy = np.sum(w * (x - wx_mean) * (y - wy_mean)) / w_sum
        var_x = np.sum(w * (x - wx_mean)**2) / w_sum
        var_y = np.sum(w * (y - wy_mean)**2) / w_sum
        if var_x == 0 or var_y == 0:
            return np.nan, np.nan
        # Weighted correlation
        corr = cov_xy / np.sqrt(var_x * var_y)
        # Compute p-value (approximate, using effective sample size)
        n_eff = len(x)  # Simplified; could use more sophisticated effective sample size
        if n_eff < 3:
            return corr, np.nan
        t_stat = corr * np.sqrt((n_eff - 2) / (1 - corr**2))
        p_value = 2 * (1 - t.cdf(np.abs(t_stat), df=n_eff-2))
        return corr, p_value
    except Exception as e:
        st.error(f"Error in weighted_correlation: {str(e)}")
        return np.nan, np.nan

def create_weighted_scatter(df, var1, var2, weights, var1_label, var2_label):
    """
    Create a weighted scatter plot with a regression line.
    
    Parameters:
    - df: pandas DataFrame containing the data.
    - var1, var2: Column names for x and y variables.
    - weights: pandas Series or array of weights for each data point.
    - var1_label, var2_label: Labels for x and y axes.
    
    Returns:
    - fig: matplotlib Figure object.
    """
    # Remove rows with NaN in var1, var2, or weights
    valid_mask = (~df[var1].isna()) & (~df[var2].isna()) & (~weights.isna())
    x = df[var1][valid_mask].values
    y = df[var2][valid_mask].values
    w = weights[valid_mask].values

    # Check if there’s enough data to plot
    if len(x) < 2 or len(y) < 2 or len(w) < 2:
        fig, ax = plt.subplots(figsize=(6, 4))
        ax.text(0.5, 0.5, "Insufficient data to plot (too many NaN values)", 
                ha='center', va='center', fontsize=10, color='red')
        ax.set_xlabel(var1_label, fontsize=8)
        ax.set_ylabel(var2_label, fontsize=8)
        plt.tight_layout()
        return fig

    # Create figure and axis objects with slightly larger size
    fig, ax = plt.subplots(figsize=(6, 4))

    # Normalize weights for visualization (between 10 and 100 for point sizes)
    w_min, w_max = w.min(), w.max()
    if w_max == w_min:  # Avoid division by zero
        w_norm = np.full_like(w, 50)  # Default size if all weights are the same
    else:
        w_norm = 10 + (90 * (w - w_min) / (w_max - w_min))

    # Compute weighted Pearson correlation coefficient
    # Center the data
    x_mean = np.average(x, weights=w)
    y_mean = np.average(y, weights=w)
    # Weighted covariance
    cov_xy = np.sum(w * (x - x_mean) * (y - y_mean)) / np.sum(w)
    # Weighted standard deviations
    var_x = np.sum(w * (x - x_mean)**2) / np.sum(w)
    var_y = np.sum(w * (y - y_mean)**2) / np.sum(w)
    if var_x == 0 or var_y == 0:
        corr = 0.0
    else:
        corr = cov_xy / np.sqrt(var_x * var_y)

    # Create scatter plot with smaller points
    scatter = ax.scatter(x, y, s=w_norm, alpha=0.5, c=w, cmap='viridis', edgecolor='none')

    # Add weighted regression line
    slope = cov_xy / var_x if var_x != 0 else 0
    intercept = y_mean - slope * x_mean
    x_range = np.array([x.min(), x.max()])
    y_range = slope * x_range + intercept
    ax.plot(x_range, y_range, "r--", alpha=0.8, label=f'Regression (r={corr:.2f})')

    # Add gridlines for better readability
    ax.grid(True, linestyle='--', alpha=0.7)

    # Wrap the axis labels to a maximum width of 50 characters
    wrapped_xlabel = textwrap.fill(var1_label, width=50)
    wrapped_ylabel = textwrap.fill(var2_label, width=50)

    # Set labels with smaller fonts
    ax.set_xlabel(wrapped_xlabel, fontsize=8)
    ax.set_ylabel(wrapped_ylabel, fontsize=8)
    
    # Add legend with smaller font
    ax.legend(fontsize=7)

    # Add colorbar with smaller label
    cbar = plt.colorbar(scatter, ax=ax)
    cbar.set_label('Weight', fontsize=7)
    cbar.ax.tick_params(labelsize=6)

    # Adjust layout to prevent label cutoff
    plt.tight_layout()

    return fig

# Function to compute BRR standard errors for correlation
def compute_brr_se_correlation(x, y, replicate_weights, corr_data, progress_bar=None):
    try:
        # Initial correlation with final student weights for reference
        main_corr, _ = weighted_correlation(x, y, corr_data['W_FSTUWT'].values)
        if np.isnan(main_corr):
            return np.nan
        
        # Compute correlations for each replicate weight
        replicate_corrs = []
        total_weights = len(replicate_weights)
        for idx, weight_col in enumerate(replicate_weights):
            # Create a DataFrame with only the weights to handle NA dropping
            data = corr_data[[weight_col, 'W_FSTUWT']].copy()
            data = data.dropna()
            if len(data) < 2:
                continue
            
            # Subset x and y to match the non-NA indices of the weights
            indices = data.index
            pos = corr_data.index.get_indexer(indices)
            pos = pos[pos != -1]  # Remove invalid indices
            if len(pos) < 2:
                continue
            x_rep = x[pos]
            y_rep = y[pos]
            w_rep = data[weight_col].values
            
            # Compute correlation with replicate weights
            rep_corr, _ = weighted_correlation(x_rep, y_rep, w_rep)
            if not np.isnan(rep_corr):
                replicate_corrs.append(rep_corr)
        
        if not replicate_corrs:
            return np.nan
        
        # Convert to array and compute standard error
        replicate_corrs = np.array(replicate_corrs)
        se = np.sqrt((1 / 80) * np.sum((replicate_corrs - main_corr) ** 2))
        return se
    except Exception as e:
        st.error(f"Error in compute_brr_se_correlation: {str(e)}")
        return np.nan

# Function to compute p-values using BRR standard errors
def compute_brr_p_value(corr, se, n):
    try:
        if np.isnan(se) or se == 0:
            return np.nan
        t_stat = corr / se
        p_value = 2 * (1 - t.cdf(np.abs(t_stat), df=n-2))
        return p_value
    except Exception as e:
        st.error(f"Error in compute_brr_p_value: {str(e)}")
        return np.nan

# Function to apply Rubin's rules for combining correlations across plausible values
def apply_rubins_rules_correlations(corrs_list, se_list, n):
    try:
        corrs_array = np.array(corrs_list)
        se_array = np.array(se_list)
        num_pvs = len(corrs_list)
        
        # Combined point estimate
        combined_corr = np.mean(corrs_array)
        
        # Within-imputation variance
        within_var = np.mean(se_array**2)
        
        # Between-imputation variance
        between_var = np.var(corrs_array, ddof=1)
        
        # Total variance
        total_var = within_var + (1 + 1/num_pvs) * between_var
        
        # Combined standard error
        combined_se = np.sqrt(total_var)
        
        # Compute p-value
        p_value = compute_brr_p_value(combined_corr, combined_se, n)
        
        return combined_corr, combined_se, p_value
    except Exception as e:
        st.error(f"Error in apply_rubins_rules_correlations: {str(e)}")
        return np.nan, np.nan, np.nan

# Function to render correlation matrix as HTML table
def render_correlation_matrix(selected_labels, corr_matrix, p_matrix):
    html_content = """
    <style>
    table, th, td {  
            font-weight: normal !important;  
    }
    
    .corr-matrix-container {
        display: inline-block;
        overflow-x: auto;
        scrollbar-width: thin;
        min-width: 0;
        margin: 20px 0;
    }
    .corr-matrix-container::-webkit-scrollbar {
        height: 8px;
    }
    .corr-matrix-container::-webkit-scrollbar-thumb {
        background-color: #888;
        border-radius: 4px;
    }
    .corr-matrix {
        table-layout: fixed;
        border-collapse: collapse;
        font-size: 14px;
        margin: 0;
    }
    .corr-matrix th, .corr-matrix td {
        border: none;
        padding: 8px;
        box-sizing: border-box;
        text-align: center;
    }
    .corr-matrix th:first-child, .corr-matrix td:first-child {
        width: 200px !important;
        text-align: left;
        white-space: normal;
        overflow-wrap: break-word;
    }
    .corr-matrix th:not(:first-child), .corr-matrix td:not(:first-child) {
        width: 100px !important;
    }
    .corr-matrix tr:nth-child(even) {
        background-color: #f9f9f9;
    }
    .corr-matrix-title {
        font-size: 16px;
        font-weight: bold;
        text-align: left;
        margin-bottom: 5px;
    }
    .corr-matrix-subtitle {
        font-size: 16px;
        font-style: italic;
        text-align: left;
        margin-bottom: 10px;
    }
    .corr-matrix-header {
        border-top: 1px solid #000;
        border-bottom: 1px solid #000;
    }
    .corr-matrix-last-row {
        border-bottom: 1px solid #000;
    }
    </style>
    <div class="corr-matrix-container">
        <div class="corr-matrix-title">Correlation Matrix</div>
        <div class="corr-matrix-subtitle">Weighted Correlations Between Selected Variables</div>
        <table class="corr-matrix">
            <tr class="corr-matrix-header">
                <th></th>
                {{header_row}}
            </tr>
            {{data_rows}}
        </table>
    </div>
    """
    # Generate header row
    header_row = ""
    for label in selected_labels:
        header_row += f"<th>{label}</th>"
    
    # Generate data rows
    data_rows = ""
    for i in range(len(selected_labels)):
        row_class = "corr-matrix-last-row" if i == len(selected_labels) - 1 else ""
        row = f'<tr class="{row_class}"><th>{selected_labels[i]}</th>'
        for j in range(len(selected_labels)):
            if i == j:
                cell_content = "1.00"
            else:
                corr = corr_matrix[i, j]
                p_value = p_matrix[i, j]
                corr_display = f"{corr:.2f}" if not np.isnan(corr) else "-"
                sig_display = "**" if p_value < 0.01 else "*" if p_value < 0.05 else "" if not np.isnan(p_value) else ""
                cell_content = f"{corr_display}{sig_display}"
            row += f"<td>{cell_content}</td>"
        row += "</tr>"
        data_rows += row
    
    full_html = html_content.replace("{{header_row}}", header_row).replace("{{data_rows}}", data_rows)
    return full_html

# Function to set cell borders in Word document
def set_cell_border(cell, top=False, bottom=False, left=False, right=False):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    # Remove existing borders
    for border in ['top', 'bottom', 'left', 'right']:
        existing_border = tcPr.find(qn(f'w:{border}'))
        if existing_border is not None:
            tcPr.remove(existing_border)
    
    # Set new borders where specified
    if top:
        border = OxmlElement('w:top')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Border width in eighths of a point
        tcPr.append(border)
    if bottom:
        border = OxmlElement('w:bottom')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        tcPr.append(border)
    if left:
        border = OxmlElement('w:left')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        tcPr.append(border)
    if right:
        border = OxmlElement('w:right')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        tcPr.append(border)

# Function to create a Word document with the APA-style table
def create_word_table(selected_labels, corr_matrix, p_matrix):
    doc = Document()
    
    # Set document to landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    
    # Add title
    title = doc.add_paragraph("Correlation Matrix")
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.runs[0].font.name = 'Times New Roman'
    title.runs[0].font.size = Pt(10)
    
    # Add subtitle
    subtitle = doc.add_paragraph("Weighted Correlations Between Selected Variables")
    subtitle.runs[0].italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle.runs[0].font.name = 'Times New Roman'
    subtitle.runs[0].font.size = Pt(10)
    
    # Create table
    table = doc.add_table(rows=1 + len(selected_labels), cols=1 + len(selected_labels))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Normal Table'  # Use a style with no borders
    
    # Set column widths
    for column in table.columns:
        for cell in column.cells:
            cell.width = Inches(0.7)  # Adjusted for landscape orientation
    
    # Add header row
    table.rows[0].cells[0].text = ""
    for idx, label in enumerate(selected_labels):
        cell = table.rows[0].cells[idx + 1]
        cell.text = label
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].bold = False
        set_cell_border(cell, top=True, bottom=True)
    # Set border for the empty cell in header row
    set_cell_border(table.rows[0].cells[0], top=True, bottom=True)
    
    # Add data rows
    for i in range(len(selected_labels)):
        row = table.rows[i + 1]
        row.cells[0].text = selected_labels[i]
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        row.cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[0].paragraphs[0].runs[0].bold = False
        for j in range(len(selected_labels)):
            cell = row.cells[j + 1]
            if i == j:
                cell_content = "1.00"
            else:
                corr = corr_matrix[i, j]
                p_value = p_matrix[i, j]
                corr_display = f"{corr:.2f}" if not np.isnan(corr) else "-"
                sig_display = "**" if p_value < 0.01 else "*" if p_value < 0.05 else "" if not np.isnan(p_value) else ""
                cell_content = f"{corr_display}{sig_display}"
            cell.text = cell_content
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if i == len(selected_labels) - 1:
                set_cell_border(cell, bottom=True)
        # Set border for the first cell in each row
        if i == len(selected_labels) - 1:
            set_cell_border(row.cells[0], bottom=True)
    
    # Save document to a BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Access data from session state
df = st.session_state.get('df', None)
variable_labels = st.session_state.get('variable_labels', {})
value_labels = st.session_state.get('value_labels', {})
visible_columns = st.session_state.get('visible_columns', [])

# Initialize session state for this page
if 'correlation_matrix_results' not in st.session_state:
    st.session_state.correlation_matrix_results = None
if 'correlation_matrix_completed' not in st.session_state:
    st.session_state.correlation_matrix_completed = False
if 'correlation_matrix_used_brr' not in st.session_state:
    st.session_state.correlation_matrix_used_brr = False

# Streamlit UI
st.title("Correlation Matrix Analysis")

if df is None or df.empty:
    st.warning("No data available. Please upload a dataset on the main page.")
else:
    # Detect all numeric variables (including all PVs for analysis)
    pv_pattern = re.compile(r'^PV([1-9]|10)(MATH|READ|SCIE)(\d*)$', re.IGNORECASE)
    weight_columns = ['W_FSTUWT'] + [f"W_FSTURWT{i}" for i in range(1, 81)] + [col for col in df.columns if 'W_FSCHWT' in col]
    excluded_variables = [
        'CYC', 'NATCEN', 'STRATUM', 'SUBNATIO', 'REGION', 'OECD', 'ADMINMODE',
        'LANGTEST_QQQ', 'LANGTEST_COG', 'LANGTEST_PAQ', 'OPTION_CT', 'OPTION_FL',
        'OPTION_ICTQ', 'OPTION_WBQ', 'OPTION_PQ', 'OPTION_TQ', 'OPTION_UH', 'BOOKID',
        'COBN_S', 'COBN_M', 'COBN_F', 'OCOD1', 'OCOD2', 'OCOD3',
        'ST001D01T', 'ST003D02T', 'ST003D03T',
        'EFFORT1', 'EFFORT2', 'PROGN', 'ISCEDP', 'SENWT', 'VER_DAT', 'TEST',
        'GRADE', 'UNIT', 'WVARSTRR',
        'PAREDINT', 'HISEI', 'HOMEPOS', 'BMMJ1', 'BFMJ2',
        'SCHOOLID', 'STUID', 'CNT', 'CNTRYID', 'CNTSCHID', 'CNTSTUID'
    ]
    item_pattern = re.compile(r'^(ST|FL|IC|WB|PA)\w{8}$', re.IGNORECASE)
    numeric_vars = [
        col for col in df.columns 
        if col not in weight_columns 
        and col not in excluded_variables
        and not item_pattern.match(col) 
        and not df[col].isna().all()
        and df[col].dtype in ['float64', 'int64']
    ]
    
    # Identify plausible value domains and regular numeric variables
    pv_domains = {}
    regular_numeric_vars = []
    for col in numeric_vars:
        pv_match = pv_pattern.match(col)
        if pv_match:
            pv_num = int(pv_match.group(1))
            domain = pv_match.group(2).upper()  # Normalize to uppercase
            if domain not in pv_domains:
                pv_domains[domain] = []
            pv_domains[domain].append(col)
        else:
            regular_numeric_vars.append(col)
    
    # Sort PVs within each domain to ensure PV1 to PV10 order
    for domain in pv_domains:
        pv_domains[domain].sort(key=lambda x: int(re.match(r'PV(\d+)', x, re.IGNORECASE).group(1)))
    
    # Check if PV domains have all 10 plausible values
    for domain, pv_list in pv_domains.items():
        if len(pv_list) != 10:
            st.warning(f"Domain {domain} has {len(pv_list)} plausible values instead of 10. Only {pv_list} will be used.")
    
    # Create domain options for selection
    domain_to_label = {
        'MATH': 'Mathematics score',
        'READ': 'Reading score',
        'SCIE': 'Science score'
    }
    domain_options = [domain_to_label.get(domain, domain) for domain in pv_domains.keys()]
    label_to_domain = {domain_to_label.get(domain, domain): domain for domain in pv_domains.keys()}
    
    # Prepare labels for regular numeric variables, using only visible ones for UI
    regular_numeric_vars = [col for col in regular_numeric_vars if col in visible_columns]
    var_labels = [variable_labels.get(col, col) for col in regular_numeric_vars]
    label_to_var = {variable_labels.get(col, col): col for col in regular_numeric_vars}
    unique_var_labels = []
    seen_labels = {}
    for col, label in zip(regular_numeric_vars, var_labels):
        if label in seen_labels:
            unique_label = f"{label} ({col})"
        else:
            unique_label = label
        seen_labels[label] = True
        unique_var_labels.append(unique_label)
        label_to_var[unique_label] = col
    
    if len(unique_var_labels) + len(domain_options) < 2:
        st.warning("At least two numeric variables or domains are required for correlation matrix analysis.")
    else:
        # Select Variables (default to MATH domain only)
        st.write("Select variables for correlation matrix (default: Mathematics score):")
        all_var_options = domain_options + unique_var_labels
        default_vars = []
        if "Mathematics score" in all_var_options:
            default_vars.append("Mathematics score")
        
        selected_var_labels = st.multiselect(
            "Variables",
            all_var_options,
            default=default_vars,
            key="matrix_vars"
        )
        
        # Split selected variables into domains and regular variables
        selected_domains = []
        selected_vars = []
        selected_labels = []
        selected_codes = []
        all_vars = []
        
        for label in selected_var_labels:
            if label in label_to_domain:
                domain = label_to_domain[label]
                selected_domains.append(domain)
                selected_labels.append(domain_to_label.get(domain, domain))
                selected_codes.append(domain)
                all_vars.append(pv_domains[domain])
            else:
                selected_vars.append(label_to_var[label])
                selected_labels.append(label)
                selected_codes.append(label_to_var[label])
                all_vars.append([label_to_var[label]])
        
        run_analysis = st.button("Run Analysis", key="run_correlation_matrix")
        
        if run_analysis and (selected_vars or selected_domains) and len(selected_var_labels) >= 2:
            try:
                if 'W_FSTUWT' not in df.columns:
                    st.error("Final student weight (W_FSTUWT) not found in the dataset.")
                else:
                    # Check for replicate weights availability
                    replicate_weight_cols = [f"W_FSTURWT{i}" for i in range(1, 81)]
                    missing_weights = [col for col in replicate_weight_cols if col not in df.columns]
                    st.session_state.correlation_matrix_used_brr = len(missing_weights) == 0  # Use BRR if no replicate weights are missing
                    
                    # Initialize correlation and p-value matrices
                    n_vars = len(selected_var_labels)
                    corr_matrix = np.ones((n_vars, n_vars))
                    p_matrix = np.zeros((n_vars, n_vars))
                    
                    # Compute the maximum number of plausible values
                    max_pvs = max([len(vars) for vars in all_vars])
                    
                    # Progress bar for the entire process
                    pv_progress = st.progress(0)
                    total_iterations = (n_vars * (n_vars - 1)) // 2 * max_pvs
                    iteration_count = 0
                    
                    # Create a placeholder for status messages
                    status_placeholder = st.empty()
                    
                    # Compute correlations for each pair of variables
                    for i in range(n_vars):
                        for j in range(i + 1, n_vars):
                            var1_list = all_vars[i]
                            var2_list = all_vars[j]
                            var1_is_pv = len(var1_list) > 1
                            var2_is_pv = len(var2_list) > 1
                            num_pvs = max(len(var1_list), len(var2_list))
                            
                            all_corrs = []
                            all_p_values = []  # Store p-values directly
                            all_se = []
                            
                            for pv_idx in range(num_pvs):
                                var1 = var1_list[pv_idx % len(var1_list)]
                                var2 = var2_list[pv_idx % len(var2_list)]
                                status_placeholder.write(f"Processing correlation between {var1} and {var2}...")
                                
                                # Create a subset DataFrame for just this pair to ensure consistent NA handling
                                pair_columns = ['W_FSTUWT']
                                if st.session_state.correlation_matrix_used_brr:
                                    pair_columns += replicate_weight_cols
                                pair_columns.append(var1)
                                pair_columns.append(var2)
                                pair_data = df[pair_columns].dropna()
                                
                                if len(pair_data) < 2:
                                    st.warning(f"Insufficient non-missing data for correlation between {var1} and {var2} (at least 2 observations required after dropping missing values).")
                                    all_corrs.append(np.nan)
                                    all_p_values.append(np.nan)
                                    all_se.append(np.nan)
                                    iteration_count += 1
                                    # Avoid division by zero in progress calculation
                                    if total_iterations > 0:
                                        pv_progress.progress(min(0.7 * (iteration_count / total_iterations), 0.7))
                                    continue
                                
                                x = pair_data[var1].values
                                y = pair_data[var2].values
                                w = pair_data['W_FSTUWT'].values
                                corr, p_value = weighted_correlation(x, y, w)
                                if st.session_state.correlation_matrix_used_brr:
                                    status_placeholder.write("Calculating standard error using replicate weights...")
                                    se = compute_brr_se_correlation(x, y, replicate_weight_cols, pair_data)
                                    p_value = compute_brr_p_value(corr, se, len(pair_data))
                                else:
                                    se = np.nan  # Not used in simple method
                                
                                all_corrs.append(corr)
                                all_p_values.append(p_value)
                                all_se.append(se)
                                
                                iteration_count += 1
                                # Avoid division by zero in progress calculation
                                if total_iterations > 0:
                                    pv_progress.progress(min(0.7 * (iteration_count / total_iterations), 0.7))
                        
                            # Combine results
                            if num_pvs == 1:
                                # For non-PV variables, use the single correlation and p-value directly
                                status_placeholder.write(f"Combining results for {selected_labels[i]} and {selected_labels[j]} (single iteration, no plausible values)...")
                                combined_corr = all_corrs[0]
                                combined_p_value = all_p_values[0]
                                combined_se = all_se[0]
                            else:
                                # For PV variables, apply Rubin's rules
                                status_placeholder.write(f"Combining results for {selected_labels[i]} and {selected_labels[j]}...")
                                combined_corr, combined_se, combined_p_value = apply_rubins_rules_correlations(all_corrs, all_se, len(pair_data))
                            
                            # Store in matrices
                            corr_matrix[i, j] = combined_corr
                            corr_matrix[j, i] = combined_corr
                            p_matrix[i, j] = combined_p_value
                            p_matrix[j, i] = combined_p_value
                    
                    # Update progress bar after correlation computation
                    pv_progress.progress(0.9)  # 90% after rendering the table
                    
                    # Store results in session state
                    st.session_state.correlation_matrix_results = {
                        'labels': selected_labels,
                        'codes': selected_codes,
                        'corr_matrix': corr_matrix,
                        'p_matrix': p_matrix
                    }
                    st.session_state.correlation_matrix_completed = True
                    
                    # Render the correlation matrix
                    status_placeholder.write("Rendering correlation matrix...")
                    table_html = render_correlation_matrix(selected_labels, corr_matrix, p_matrix)
                    # Calculate height based on number of variables (approximately 90px per row plus some padding)
                    matrix_height = len(selected_labels) * 90 + 30
                    components.html(table_html, height=matrix_height, scrolling=True)
                    status_placeholder.write("Correlation matrix analysis completed.")
                    
                    # Provide download button for Word document (moved here)
                    doc_buffer = create_word_table(selected_labels, corr_matrix, p_matrix)
                    st.download_button(
                        label="Download Table as Word Document",
                        data=doc_buffer,
                        file_name="Correlation_Matrix_Table.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    # Generate scatter plots for all pairs in a 2-column grid
                    st.header("Scatter plots for all variable pairs:")
                    variable_pairs = list(itertools.combinations(range(n_vars), 2))
                    status_placeholder.write(f"Generating scatter plots for {len(variable_pairs)} pairs of variables.")
                    
                    for idx in range(0, len(variable_pairs), 2):
                        # Create a row with 2 columns
                        cols = st.columns(2)
                        
                        # First plot in the row
                        i, j = variable_pairs[idx]
                        var1_label = selected_labels[i]
                        var2_label = selected_labels[j]
                        var1_code = selected_codes[i]
                        var2_code = selected_codes[j]
                        
                        # Handle plausible value domains
                        if var1_code in pv_domains:
                            var1_data = df[pv_domains[var1_code][0]]
                        else:
                            var1_data = df[var1_code]
                        
                        if var2_code in pv_domains:
                            var2_data = df[pv_domains[var2_code][0]]
                        else:
                            var2_data = df[var2_code]
                        
                        # Create plot dataframe step-by-step
                        data_dict = {
                            var1_code: var1_data,
                            var2_code: var2_data,
                            'weights': df['W_FSTUWT']
                        }
                        plot_df = pd.DataFrame(data_dict)
                        plot_df = plot_df.dropna()
                        
                        # Create scatter plot
                        fig1 = create_weighted_scatter(plot_df, var1_code, var2_code, 
                                                      plot_df['weights'], var1_label, var2_label)
                        with cols[0]:
                            st.pyplot(fig1)
                        
                        # Second plot in the row (if available)
                        if idx + 1 < len(variable_pairs):
                            i, j = variable_pairs[idx + 1]
                            var1_label = selected_labels[i]
                            var2_label = selected_labels[j]
                            var1_code = selected_codes[i]
                            var2_code = selected_codes[j]
                            
                            # Handle plausible value domains
                            if var1_code in pv_domains:
                                var1_data = df[pv_domains[var1_code][0]]
                            else:
                                var1_data = df[var1_code]
                            
                            if var2_code in pv_domains:
                                var2_data = df[pv_domains[var2_code][0]]
                            else:
                                var2_data = df[var2_code]
                            
                            # Create plot dataframe step-by-step
                            data_dict = {
                                var1_code: var1_data,
                                var2_code: var2_data,
                                'weights': df['W_FSTUWT']
                            }
                            plot_df = pd.DataFrame(data_dict)
                            plot_df = plot_df.dropna()
                            
                            # Create scatter plot
                            fig2 = create_weighted_scatter(plot_df, var1_code, var2_code, 
                                                          plot_df['weights'], var1_label, var2_label)
                            with cols[1]:
                                st.pyplot(fig2)
                    
                    # Update progress bar to 100% after scatter plots
                    pv_progress.progress(1.0)  # 100% after scatter plots are generated
                    status_placeholder.empty()  # Clear the placeholder after completion
            except Exception as e:
                st.error(f"Error computing correlation matrix: {str(e)}")
                st.session_state.correlation_matrix_completed = False
                st.session_state.correlation_matrix_results = None
        elif st.session_state.correlation_matrix_results and st.session_state.correlation_matrix_completed:
            if selected_var_labels:
                results = st.session_state.correlation_matrix_results
                table_html = render_correlation_matrix(
                    results['labels'],
                    results['corr_matrix'],
                    results['p_matrix']
                )
                components.html(table_html, height=400, scrolling=True)
                st.write("Correlation matrix analysis completed.")
                
                # Provide download button for Word document (moved here)
                doc_buffer = create_word_table(
                    results['labels'],
                    results['corr_matrix'],
                    results['p_matrix']
                )
                st.download_button(
                    label="Download Table as Word Document",
                    data=doc_buffer,
                    file_name="Correlation_Matrix_Table.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.write("Please select at least two variables to compute the correlation matrix.")
        else:
            st.write("Please select at least two variables and click 'Run Analysis' to compute the correlation matrix.")

# Instructions section
st.header("Instructions")
st.markdown("""
- **Select Variables**: Choose two or more variables (domains or numeric variables) from the dropdown menus. Plausible value domains (e.g., Mathematics score, Reading score) will use all 10 plausible values for analysis.
- **Run Analysis**: Click "Run Analysis" to perform the weighted correlation matrix analysis. Analyses involving plausible values will be combined using Rubin's rules.
- **View Results**: Results are displayed in an APA-style table with correlations rounded to 2 decimal places. You can download the table as a Word document using the download button.
- **Navigate**: Use the sidebar to switch between different analysis types or return to the main page to upload a new dataset.
""")