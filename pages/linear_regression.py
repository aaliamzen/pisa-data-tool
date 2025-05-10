import streamlit as st
import pandas as pd
import numpy as np
import re
import statsmodels.api as sm
import streamlit.components.v1 as components
from scipy.stats import t, anderson, probplot
from statsmodels.stats.diagnostic import het_breuschpagan
from statsmodels.stats.outliers_influence import variance_inflation_factor
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import seaborn as sns
import base64

# Add logo that persists across all pages
try:
    st.logo("assets/logo.png")  # Replace with the path to your logo file, e.g., "assets/logo.png"
except Exception as e:
    st.error(f"Failed to load logo: {e}")

# Streamlit app configuration
st.set_page_config(page_title="Linear Regression - PISA Data Exploration Tool", layout="wide")

# Function to compute weighted standard deviation
def weighted_std(x, w):
    try:
        # Remove NaN values
        mask = ~(np.isnan(x) | np.isnan(w))
        x = x[mask]
        w = w[mask]
        if len(x) < 2 or np.sum(w) == 0:
            return np.nan
        
        # Weighted mean
        w_sum = np.sum(w)
        w_mean = np.sum(w * x) / w_sum
        
        # Weighted variance
        variance = np.sum(w * (x - w_mean)**2) / w_sum
        
        # Weighted standard deviation
        return np.sqrt(variance)
    except Exception as e:
        st.error(f"Error in weighted_std: {str(e)}")
        return np.nan

# Function to compute weighted OLS regression for a single set of variables
def weighted_ols_regression(x, y, w, var_names):
    try:
        # Remove NaN values
        mask = ~(np.isnan(y) | np.isnan(w) | np.any(np.isnan(x), axis=1))
        x = x[mask]
        y = y[mask]
        w = w[mask]
        if len(x) < 2:
            return [(var, np.nan, np.nan, np.nan, np.nan) for var in ['Intercept'] + var_names], np.nan, np.nan, {}
        
        # Compute weighted standard deviations for standardization
        sd_y = weighted_std(y, w)
        sd_x = [weighted_std(x[:, i], w) for i in range(x.shape[1])]
        
        # Validate standard deviations
        if np.isnan(sd_y) or sd_y == 0:
            st.warning("Standard deviation of the dependent variable is zero or NaN. Standardized coefficients (β) cannot be computed.")
        
        for i, sd in enumerate(sd_x):
            if np.isnan(sd) or sd == 0:
                st.warning(f"Standard deviation of predictor {var_names[i]} is zero or NaN. Standardized coefficient (β) for this predictor will be NaN.")
        
        # Add constant for intercept
        X = sm.add_constant(x)
        
        # Fit weighted OLS model
        model = sm.WLS(y, X, weights=w).fit()
        
        # Extract coefficients, standard errors, p-values, and compute standardized coefficients
        results = []
        for idx, var in enumerate(['Intercept'] + var_names):
            coef = model.params[idx]
            se = model.bse[idx]
            p_value = model.pvalues[idx]
            # Compute standardized coefficient (β)
            if idx == 0:  # Intercept
                std_coef = np.nan  # Standardized intercept is not typically meaningful
            else:
                if np.isnan(sd_y) or sd_y == 0 or np.isnan(sd_x[idx-1]) or sd_x[idx-1] == 0:
                    std_coef = np.nan
                else:
                    std_coef = coef * (sd_x[idx-1] / sd_y)
            results.append((var, coef, std_coef, se, p_value))
        
        # Get R-squared and adjusted R-squared values
        r_squared = model.rsquared if hasattr(model, 'rsquared') else np.nan
        r_squared_adj = model.rsquared_adj if hasattr(model, 'rsquared_adj') else np.nan
        
        # Compute model diagnostics
        diagnostics = {}
        
        # F-statistic and p-value
        f_stat = model.fvalue if hasattr(model, 'fvalue') else np.nan
        f_pvalue = model.f_pvalue if hasattr(model, 'f_pvalue') else np.nan
        df_model = model.df_model if hasattr(model, 'df_model') else np.nan
        df_resid = model.df_resid if hasattr(model, 'df_resid') else np.nan
        diagnostics['f_stat'] = (f_stat, f_pvalue, df_model, df_resid)
        
        # Normality of residuals (Anderson-Darling test)
        if len(model.resid) > 3:  # Ensure enough data for the test
            ad_result = anderson(model.resid, dist='norm')
            ad_stat = ad_result.statistic
            # Check if the test statistic exceeds the critical value at 5% significance level
            critical_value_5 = ad_result.critical_values[2]  # Index 2 corresponds to 5% significance
            ad_reject_normality = ad_stat > critical_value_5
        else:
            ad_stat, ad_reject_normality = np.nan, False
        diagnostics['anderson_darling'] = (ad_stat, ad_reject_normality)
        
        # Homoscedasticity (Breusch-Pagan test)
        if len(x) > len(var_names) + 1:  # Ensure enough data for the test
            bp_lm_stat, bp_pvalue, _, _ = het_breuschpagan(model.resid, X)
        else:
            bp_lm_stat, bp_pvalue = np.nan, np.nan
        diagnostics['breusch_pagan'] = (bp_lm_stat, bp_pvalue)
        
        # Multicollinearity (VIF)
        vif_values = []
        if x.shape[1] > 0:  # Ensure there are predictors to compute VIF
            for i in range(x.shape[1]):
                vif = variance_inflation_factor(X, i + 1)  # Skip the intercept (column 0)
                vif_values.append(vif)
        max_vif = max(vif_values) if vif_values else np.nan
        diagnostics['max_vif'] = max_vif
        
        return results, r_squared, r_squared_adj, diagnostics
    except Exception as e:
        st.error(f"Error in weighted_ols_regression: {str(e)}")
        return [(var, np.nan, np.nan, np.nan, np.nan) for var in ['Intercept'] + var_names], np.nan, np.nan, {}

# Function to compute BRR standard errors for regression coefficients
def compute_brr_se_regression(x, y, replicate_weights, reg_data, progress_bar=None, var_names=None):
    try:
        # Initial regression with final student weights for reference
        main_results, main_r_squared, main_r_squared_adj, main_diagnostics = weighted_ols_regression(x, y, reg_data['W_FSTUWT'].values, var_names)
        main_coefs = [result[1] for result in main_results]  # Unstandardized coefficients
        main_std_coefs = [result[2] for result in main_results]  # Standardized coefficients
        if all(np.isnan(main_coefs)):
            return ([np.nan] * len(main_coefs), [np.nan] * len(main_std_coefs))
        
        # Compute coefficients for each replicate weight
        replicate_coefs = {idx: [] for idx in range(len(main_coefs))}
        replicate_std_coefs = {idx: [] for idx in range(len(main_std_coefs))}
        total_weights = len(replicate_weights)
        for idx, weight_col in enumerate(replicate_weights):
            # Create a DataFrame with only the weights to handle NA dropping
            data = reg_data[[weight_col, 'W_FSTUWT']].copy()
            data = data.dropna()
            if len(data) < 2:
                continue
            
            # Subset x and y to match the non-NA indices of the weights
            indices = data.index
            pos = reg_data.index.get_indexer(indices)
            pos = pos[pos != -1]  # Remove invalid indices
            if len(pos) < 2:
                continue
            x_rep = x[pos]
            y_rep = y[pos]
            w_rep = data[weight_col].values
            
            # Compute regression with replicate weights
            rep_results, _, _, _ = weighted_ols_regression(x_rep, y_rep, w_rep, var_names)
            for coef_idx, (var, coef, std_coef, _, _) in enumerate(rep_results):
                if not np.isnan(coef):
                    replicate_coefs[coef_idx].append(coef)
                if not np.isnan(std_coef):
                    replicate_std_coefs[coef_idx].append(std_coef)
        
        # Compute BRR standard errors for each coefficient (unstandardized and standardized)
        brr_se = []
        brr_se_std = []
        for idx, (var, main_coef, main_std_coef, _, _) in enumerate(main_results):
            # Unstandardized coefficient SE
            if not replicate_coefs[idx]:
                brr_se.append(np.nan)
            else:
                replicate_coefs_array = np.array(replicate_coefs[idx])
                se = np.sqrt((1 / 80) * np.sum((replicate_coefs_array - main_coef) ** 2))
                brr_se.append(se)
            
            # Standardized coefficient SE
            if not replicate_std_coefs[idx]:
                brr_se_std.append(np.nan)
            else:
                replicate_std_coefs_array = np.array(replicate_std_coefs[idx])
                se_std = np.sqrt((1 / 80) * np.sum((replicate_std_coefs_array - main_std_coef) ** 2))
                brr_se_std.append(se_std)
        
        return (brr_se, brr_se_std)
    except Exception as e:
        st.error(f"Error in compute_brr_se_regression: {str(e)}")
        return ([np.nan] * (len(var_names) + 1), [np.nan] * (len(var_names) + 1))

# Function to compute p-values using BRR standard errors
def compute_brr_p_value(coef, se, n):
    try:
        if np.isnan(se) or se == 0 or np.isnan(coef):
            return np.nan
        t_stat = coef / se
        p_value = 2 * (1 - t.cdf(np.abs(t_stat), df=n-2))
        return p_value
    except Exception as e:
        st.error(f"Error in compute_brr_p_value: {str(e)}")
        return np.nan

# Function to apply Rubin's rules for combining regression results across plausible values
def apply_rubins_rules_regression(all_results, n, var_names):
    try:
        # Separate results, R-squared, adjusted R-squared, and diagnostics
        results_list = [result[0] for result in all_results]  # List of (var, coef, std_coef, se, p_value) tuples
        r_squared_list = [result[1] for result in all_results]  # List of R-squared values
        r_squared_adj_list = [result[2] for result in all_results]  # List of adjusted R-squared values
        diagnostics_list = [result[3] for result in all_results]  # List of diagnostics dictionaries
        
        # Transpose results to group by variable: list of lists where each inner list contains results for one variable across PVs
        combined_results = []
        var_results = list(zip(*results_list))  # One list per variable (including Intercept)
        
        for idx, var_result in enumerate(var_results):
            # Determine the variable name
            var = 'Intercept' if idx == 0 else var_names[idx - 1]
            
            coefs_list = [result[1] for result in var_result]  # Unstandardized coefficients
            std_coefs_list = [result[2] for result in var_result]  # Standardized coefficients
            se_list = [result[3] for result in var_result]
            
            # Skip if all coefficients are NaN
            if np.all(np.isnan(coefs_list)):
                combined_results.append((var, np.nan, np.nan, np.nan, np.nan))
                continue
            
            num_pvs = len([c for c in coefs_list if not np.isnan(c)])
            if num_pvs == 0:
                combined_results.append((var, np.nan, np.nan, np.nan, np.nan))
                continue
            
            # Combined point estimate (unstandardized coefficient)
            combined_coef = np.nanmean(coefs_list)
            
            # Combined point estimate (standardized coefficient)
            # Check if there are any non-NaN standardized coefficients to avoid warning
            std_coefs_array = np.array(std_coefs_list)
            if np.all(np.isnan(std_coefs_array)):
                combined_std_coef = np.nan
            else:
                combined_std_coef = np.nanmean(std_coefs_array)
            
            # Within-imputation variance
            within_var = np.nanmean(np.array(se_list)**2)
            
            # Between-imputation variance (unstandardized coefficient)
            between_var = np.nanvar(coefs_list, ddof=1)
            
            # Total variance (unstandardized coefficient)
            total_var = within_var + (1 + 1/num_pvs) * between_var
            
            # Combined standard error (unstandardized coefficient)
            combined_se = np.sqrt(total_var)
            
            # Compute p-value
            p_value = compute_brr_p_value(combined_coef, combined_se, n)
            
            combined_results.append((var, combined_coef, combined_std_coef, combined_se, p_value))
        
        # Average R-squared and adjusted R-squared across PVs
        r_squared = np.nanmean(r_squared_list)
        r_squared_adj = np.nanmean(r_squared_adj_list)
        
        # Combine diagnostics across PVs
        combined_diagnostics = {}
        
        # F-statistic: average the statistic and p-value, use the last df values
        f_stats = [d.get('f_stat', (np.nan, np.nan, np.nan, np.nan))[0] for d in diagnostics_list]
        f_pvalues = [d.get('f_stat', (np.nan, np.nan, np.nan, np.nan))[1] for d in diagnostics_list]
        df_model = diagnostics_list[-1].get('f_stat', (np.nan, np.nan, np.nan, np.nan))[2]
        df_resid = diagnostics_list[-1].get('f_stat', (np.nan, np.nan, np.nan, np.nan))[3]
        combined_diagnostics['f_stat'] = (np.nanmean(f_stats), np.nanmean(f_pvalues), df_model, df_resid)
        
        # Anderson-Darling: average the statistic, combine rejection decisions with logical OR
        ad_stats = [d.get('anderson_darling', (np.nan, False))[0] for d in diagnostics_list]
        ad_reject = any(d.get('anderson_darling', (np.nan, False))[1] for d in diagnostics_list)
        combined_diagnostics['anderson_darling'] = (np.nanmean(ad_stats), ad_reject)
        
        # Breusch-Pagan: average the statistic and p-value
        bp_lm_stats = [d.get('breusch_pagan', (np.nan, np.nan))[0] for d in diagnostics_list]
        bp_pvalues = [d.get('breusch_pagan', (np.nan, np.nan))[1] for d in diagnostics_list]
        combined_diagnostics['breusch_pagan'] = (np.nanmean(bp_lm_stats), np.nanmean(bp_pvalues))
        
        # Max VIF: take the maximum across PVs
        max_vifs = [d.get('max_vif', np.nan) for d in diagnostics_list]
        combined_diagnostics['max_vif'] = np.nanmax(max_vifs)
        
        return combined_results, r_squared, r_squared_adj, combined_diagnostics
    except Exception as e:
        st.error(f"Error in apply_rubins_rules_regression: {str(e)}")
        return [(var, np.nan, np.nan, np.nan, np.nan) for var in ['Intercept'] + var_names], np.nan, np.nan, {}

# Function to plot to base64
def plot_to_base64():
    buf = BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return base64.b64encode(buf.read()).decode('utf-8')

# Function to compute linear regression with PVs and BRR
def compute_linear_regression_with_pvs(df, dependent_var, independent_vars, weights, replicate_weights, use_brr, var_labels, status_placeholder=None):
    try:
        # Determine if dependent or independent variables are PV domains
        dep_is_pv = dependent_var.startswith("PV")
        indep_is_pv = [var.startswith("PV") for var in independent_vars]
        
        # Map independent variable codes to their labels
        code_to_label = {code: label for label, code in var_labels.items()}
        indep_var_labels = []
        for var in independent_vars:
            if var.startswith("PV"):
                # For PV domains, use the domain label (e.g., "Mathematics score")
                pv_match = pv_pattern.match(var)
                domain = pv_match.group(2).upper()
                domain_label = domain_to_label.get(domain, domain)
                indep_var_labels.append(domain_label)
            else:
                # For regular variables, use the label from code_to_label
                indep_var_labels.append(code_to_label.get(var, var))
        
        # If neither dependent nor independent vars are PVs, run a single regression
        if not dep_is_pv and not any(indep_is_pv):
            data = df[[dependent_var] + independent_vars + ['W_FSTUWT'] + replicate_weights].dropna()
            if len(data) < 2:
                raise ValueError("Insufficient non-missing data for regression.")
            
            X = data[independent_vars].values
            y = data[dependent_var].values
            w = data['W_FSTUWT'].values
            
            # Run regression
            results, r_squared, r_squared_adj, diagnostics = weighted_ols_regression(X, y, w, independent_vars)
            if use_brr:
                if status_placeholder:
                    status_placeholder.write("Calculating standard errors using replicate weights...")
                brr_se, brr_se_std = compute_brr_se_regression(X, y, replicate_weights, data, independent_vars)
                # Update results with BRR standard errors and recompute p-values
                updated_results = []
                for idx, (var, coef, std_coef, _, _) in enumerate(results):
                    se = brr_se[idx]
                    p_value = compute_brr_p_value(coef, se, len(data))
                    # Replace variable code with label
                    if idx == 0:  # Intercept
                        updated_var = var
                    else:
                        updated_var = indep_var_labels[idx - 1]
                    updated_results.append((updated_var, coef, std_coef, se, p_value))
                results = updated_results
            else:
                # Replace variable codes with labels in the results
                updated_results = []
                for idx, (var, coef, std_coef, se, p_value) in enumerate(results):
                    if idx == 0:  # Intercept
                        updated_var = var
                    else:
                        updated_var = indep_var_labels[idx - 1]
                    updated_results.append((updated_var, coef, std_coef, se, p_value))
                results = updated_results
            
            # Generate visualizations using the final model
            visualizations = {}
            X = sm.add_constant(X)
            fitted_values = np.dot(X, [r[1] for r in results])  # Compute fitted values using combined coefficients
            residuals = y - fitted_values
            
            # Q-Q Plot
            plt.figure(figsize=(6, 4))
            probplot(residuals, dist="norm", plot=plt)
            plt.title("Q-Q Plot of Residuals")
            visualizations['qq_plot'] = plot_to_base64()
            
            # Residuals vs. Fitted Values Plot
            plt.figure(figsize=(6, 4))
            plt.scatter(fitted_values, residuals, alpha=0.5)
            plt.axhline(y=0, color='r', linestyle='--')
            plt.xlabel("Fitted Values")
            plt.ylabel("Residuals")
            plt.title("Residuals vs. Fitted Values")
            visualizations['resid_vs_fitted'] = plot_to_base64()
            
            # Residuals Histogram
            plt.figure(figsize=(6, 4))
            sns.histplot(residuals, kde=True, stat="density")
            plt.xlabel("Residuals")
            plt.title("Histogram of Residuals")
            visualizations['resid_histogram'] = plot_to_base64()
            
            return results, r_squared, r_squared_adj, diagnostics, visualizations
        
        # If either dependent or independent vars are PVs, handle PV analysis
        pv_domain = None
        if dep_is_pv:
            pv_match = pv_pattern.match(dependent_var)
            pv_domain = pv_match.group(2).upper()
            pv_nums = list(range(1, 11))  # PV1 to PV10
            pv_vars = [f"PV{i}{pv_domain}" for i in pv_nums]
        else:
            # Check if any independent variable is a PV
            for var in independent_vars:
                pv_match = pv_pattern.match(var)
                if pv_match:
                    pv_domain = pv_match.group(2).upper()
                    pv_nums = list(range(1, 11))
                    pv_vars = [f"PV{i}{pv_domain}" for i in pv_nums]
                    break
        
        if not pv_domain:
            raise ValueError("No plausible value domain identified.")
        
        # Progress bar for PV iterations
        pv_progress = st.progress(0)
        total_iterations = len(pv_vars)
        iteration_count = 0
        
        all_results = []
        for pv_idx in pv_nums:
            pv_var = f"PV{pv_idx}{pv_domain}"
            if pv_var not in df.columns:
                continue
            
            if status_placeholder:
                status_placeholder.write(f"Processing regression with {pv_var}...")
            
            # Prepare variables for this PV iteration
            if dep_is_pv:
                dep_var = pv_var
                indep_vars = independent_vars
            else:
                dep_var = dependent_var
                indep_vars = []
                for var in independent_vars:
                    if var.startswith("PV"):
                        pv_match = pv_pattern.match(var)
                        if pv_match.group(2).upper() == pv_domain:
                            indep_vars.append(f"PV{pv_idx}{pv_domain}")
                        else:
                            indep_vars.append(var)
                    else:
                        indep_vars.append(var)
            
            # Create a subset DataFrame for this iteration
            columns = [dep_var] + indep_vars + ['W_FSTUWT']
            if use_brr:
                columns += replicate_weights
            data = df[columns].dropna()
            if len(data) < 2:
                st.warning(f"Insufficient non-missing data for regression with {pv_var}.")
                iteration_count += 1
                pv_progress.progress(iteration_count / total_iterations)
                continue
            
            X = data[indep_vars].values
            y = data[dep_var].values
            w = data['W_FSTUWT'].values
            
            # Run regression
            results, r_squared, r_squared_adj, diagnostics = weighted_ols_regression(X, y, w, indep_vars)
            if use_brr:
                if status_placeholder:
                    status_placeholder.write("Calculating standard errors using replicate weights...")
                brr_se, brr_se_std = compute_brr_se_regression(X, y, replicate_weights, data, None, indep_vars)
                # Update results with BRR standard errors and recompute p-values
                updated_results = []
                for idx, (var, coef, std_coef, _, _) in enumerate(results):
                    se = brr_se[idx]
                    p_value = compute_brr_p_value(coef, se, len(data))
                    # Replace variable code with label
                    if idx == 0:  # Intercept
                        updated_var = var
                    else:
                        updated_var = indep_var_labels[idx - 1]
                    updated_results.append((updated_var, coef, std_coef, se, p_value))
                results = updated_results
            else:
                # Replace variable codes with labels in the results
                updated_results = []
                for idx, (var, coef, std_coef, se, p_value) in enumerate(results):
                    if idx == 0:  # Intercept
                        updated_var = var
                    else:
                        updated_var = indep_var_labels[idx - 1]
                    updated_results.append((updated_var, coef, std_coef, se, p_value))
                results = updated_results
            
            all_results.append((results, r_squared, r_squared_adj, diagnostics))
            
            iteration_count += 1
            pv_progress.progress(iteration_count / total_iterations)
        
        # Combine results using Rubin's rules
        if not all_results:
            raise ValueError("No valid regression results computed for any plausible values.")
        
        if status_placeholder:
            status_placeholder.write("Combining results across plausible values...")
        combined_results, combined_r_squared, combined_r_squared_adj, combined_diagnostics = apply_rubins_rules_regression(all_results, len(data), indep_var_labels)
        
        # Generate visualizations using the final combined model
        visualizations = {}
        # Recompute residuals and fitted values using the combined coefficients
        data = df[[dependent_var] + independent_vars + ['W_FSTUWT']].dropna()
        X = data[independent_vars].values
        y = data[dependent_var].values
        w = data['W_FSTUWT'].values
        X = sm.add_constant(X)
        fitted_values = np.dot(X, [r[1] for r in combined_results])
        residuals = y - fitted_values
        
        # Q-Q Plot
        plt.figure(figsize=(6, 4))
        probplot(residuals, dist="norm", plot=plt)
        plt.title("Q-Q Plot of Residuals")
        visualizations['qq_plot'] = plot_to_base64()
        
        # Residuals vs. Fitted Values Plot
        plt.figure(figsize=(6, 4))
        plt.scatter(fitted_values, residuals, alpha=0.5)
        plt.axhline(y=0, color='r', linestyle='--')
        plt.xlabel("Fitted Values")
        plt.ylabel("Residuals")
        plt.title("Residuals vs. Fitted Values")
        visualizations['resid_vs_fitted'] = plot_to_base64()
        
        # Residuals Histogram
        plt.figure(figsize=(6, 4))
        sns.histplot(residuals, kde=True, stat="density")
        plt.xlabel("Residuals")
        plt.title("Histogram of Residuals")
        visualizations['resid_histogram'] = plot_to_base64()
        
        return combined_results, combined_r_squared, combined_r_squared_adj, combined_diagnostics, visualizations
    except Exception as e:
        st.error(f"Error in compute_linear_regression_with_pvs: {str(e)}")
        return [(var, np.nan, np.nan, np.nan, np.nan) for var in ['Intercept'] + independent_vars], np.nan, np.nan, {}, {}

# Function to render regression table as HTML
def render_regression_table(dependent_var_label, results, r_squared, r_squared_adj, diagnostics):
    html_content = """
    <style>
    .reg-table-container {
        display: inline-block;
        overflow-x: auto;
        scrollbar-width: thin;
        min-width: 0;
        margin: 20px 0;
    }
    .reg-table-container::-webkit-scrollbar {
        height: 8px;
    }
    .reg-table-container::-webkit-scrollbar-thumb {
        background-color: #888;
        border-radius: 4px;
    }
    .reg-table {
        table-layout: fixed;
        border-collapse: collapse;
        font-size: 14px;
        margin: 0;
    }
    .reg-table th, .reg-table td {
        border: none;
        padding: 8px;
        box-sizing: border-box;
        text-align: center;
        font-weight: normal;  /* Remove bold styling */
    }
    .reg-table th:first-child, .reg-table td:first-child {
        width: 200px !important;
        text-align: left;
        white-space: normal;
        overflow-wrap: break-word;
    }
    .reg-table th:not(:first-child), .reg-table td:not(:first-child) {
        width: 100px !important;
    }
    .reg-table tr:nth-child(even) {
        background-color: #f9f9f9;
    }
    .reg-table-title {
        font-size: 16px;
        font-weight: bold;
        text-align: left;
        margin-bottom: 5px;
    }
    .reg-table-subtitle {
        font-size: 16px;
        font-style: italic;
        text-align: left;
        margin-bottom: 10px;
    }
    .reg-table-header {
        border-top: 1px solid #000;
        border-bottom: 1px solid #000;
    }
    .reg-table-last-row {
        border-bottom: 1px solid #000;
    }
    .reg-table-note {
        font-size: 14px;
        text-align: left;
        margin-top: 5px;
    }
    </style>
    <div class="reg-table-container">
        <div class="reg-table-title">Table 1</div>
        <div class="reg-table-subtitle">Weighted Linear Regression Results for Dependent Variable: {{dependent_var}}</div>
        <table class="reg-table">
            <tr class="reg-table-header">
                <th>Variable</th>
                <th><i>B</i></th>
                <th><i>β</i></th>
                <th><i>SE</i></th>
                <th><i>p</i></th>
            </tr>
            {{data_rows}}
        </table>
        <div class="reg-table-note"><i>Note.</i> <i>R²</i> = {{r_squared}}, Adjusted <i>R²</i> = {{r_squared_adj}}</div>
        <div class="reg-table-note">Model: <i>F</i>({{df_model}}, {{df_resid}}) = {{f_stat}}, <i>p</i> = {{f_pvalue}}</div>
        <div class="reg-table-note">Assumptions: Anderson-Darling: <i>A²</i> = {{ad_stat}}, Normality Rejected at 5% = {{ad_reject}}; Breusch-Pagan: <i>LM</i> = {{bp_lm_stat}}, <i>p</i> = {{bp_pvalue}}; Max <i>VIF</i> = {{max_vif}}</div>
    </div>
    """
    data_rows = ""
    for idx, (var, coef, std_coef, se, p_value) in enumerate(results):
        coef_display = f"{coef:.2f}" if not np.isnan(coef) else "-"
        std_coef_display = f"{std_coef:.2f}" if not np.isnan(std_coef) else "-"
        se_display = f"{se:.2f}" if not np.isnan(se) else "-"
        p_display = "< .001" if p_value < 0.001 else f"{p_value:.2f}" if not np.isnan(p_value) else "-"
        sig_display = "**" if p_value < 0.01 else "*" if p_value < 0.05 else "" if not np.isnan(p_value) else ""
        row_class = "reg-table-last-row" if idx == len(results) - 1 else ""
        row = f"""
        <tr class="{row_class}">
            <th>{var}</th>
            <td>{coef_display}{sig_display}</td>
            <td>{std_coef_display}</td>
            <td>{se_display}</td>
            <td>{p_display}</td>
        </tr>
        """
        data_rows += row
    
    r_squared_display = f"{r_squared:.3f}" if not np.isnan(r_squared) else "-"
    r_squared_adj_display = f"{r_squared_adj:.3f}" if not np.isnan(r_squared_adj) else "-"
    
    # Extract diagnostics
    f_stat, f_pvalue, df_model, df_resid = diagnostics.get('f_stat', (np.nan, np.nan, np.nan, np.nan))
    f_stat_display = f"{f_stat:.2f}" if not np.isnan(f_stat) else "-"
    f_pvalue_display = "< .001" if f_pvalue < 0.001 else f"{f_pvalue:.3f}" if not np.isnan(f_pvalue) else "-"
    df_model_display = f"{int(df_model)}" if not np.isnan(df_model) else "-"
    df_resid_display = f"{int(df_resid)}" if not np.isnan(df_resid) else "-"
    
    ad_stat, ad_reject = diagnostics.get('anderson_darling', (np.nan, False))
    ad_stat_display = f"{ad_stat:.2f}" if not np.isnan(ad_stat) else "-"
    ad_reject_display = "Yes" if ad_reject else "No"
    
    bp_lm_stat, bp_pvalue = diagnostics.get('breusch_pagan', (np.nan, np.nan))
    bp_lm_stat_display = f"{bp_lm_stat:.2f}" if not np.isnan(bp_lm_stat) else "-"
    bp_pvalue_display = "< .001" if bp_pvalue < 0.001 else f"{bp_pvalue:.3f}" if not np.isnan(bp_pvalue) else "-"
    
    max_vif = diagnostics.get('max_vif', np.nan)
    max_vif_display = f"{max_vif:.2f}" if not np.isnan(max_vif) else "-"
    
    full_html = html_content.replace("{{dependent_var}}", dependent_var_label).replace("{{data_rows}}", data_rows).replace("{{r_squared}}", r_squared_display).replace("{{r_squared_adj}}", r_squared_adj_display).replace("{{f_stat}}", f_stat_display).replace("{{f_pvalue}}", f_pvalue_display).replace("{{df_model}}", df_model_display).replace("{{df_resid}}", df_resid_display).replace("{{ad_stat}}", ad_stat_display).replace("{{ad_reject}}", ad_reject_display).replace("{{bp_lm_stat}}", bp_lm_stat_display).replace("{{bp_pvalue}}", bp_pvalue_display).replace("{{max_vif}}", max_vif_display)
    
    return full_html

# Function to set cell borders in Word document
def set_cell_border(cell, top=False, bottom=False, left=False, right=False):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    # Remove existing borders
    for border in ['top', 'bottom', 'left', 'right']:
        existing_border = tcPr.find(qn(f'w:{border}'))
        if existing_border is not None:
            tcPr.remove(existing_border)
    
    # Set new borders where specified
    if top:
        border = OxmlElement('w:top')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Border width in eighths of a point
        tcPr.append(border)
    if bottom:
        border = OxmlElement('w:bottom')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        tcPr.append(border)
    if left:
        border = OxmlElement('w:left')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        tcPr.append(border)
    if right:
        border = OxmlElement('w:right')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        tcPr.append(border)

# Function to create a Word document with the APA-style table
def create_word_table(dependent_var_label, results, r_squared, r_squared_adj, diagnostics, visualizations):
    doc = Document()
    
    # Set document to landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    
    # Add title
    title = doc.add_paragraph("Table 1")
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.runs[0].font.name = 'Times New Roman'
    title.runs[0].font.size = Pt(10)
    
    # Add subtitle
    subtitle = doc.add_paragraph(f"Weighted Linear Regression Results for Dependent Variable: {dependent_var_label}")
    subtitle.runs[0].italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle.runs[0].font.name = 'Times New Roman'
    subtitle.runs[0].font.size = Pt(10)
    
    # Create table
    table = doc.add_table(rows=1 + len(results), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Normal Table'  # Use a style with no borders
    
    # Set column widths
    for column in table.columns:
        for cell in column.cells:
            cell.width = Inches(1.2)  # Adjusted for landscape orientation and additional column
    
    # Add header row
    headers = ["Variable", "B", "β", "SE", "p"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        p = cell.paragraphs[0]
        if header in ["B", "β", "SE", "p"]:
            # Create a run for the italicized text
            run = p.add_run(header)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.italic = True
        else:
            cell.text = header
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = False
        set_cell_border(cell, top=True, bottom=True)
    
    # Add data rows
    for idx, (var, coef, std_coef, se, p_value) in enumerate(results):
        row = table.rows[idx + 1]
        row.cells[0].text = var
        coef_display = f"{coef:.2f}" if not np.isnan(coef) else "-"
        std_coef_display = f"{std_coef:.2f}" if not np.isnan(std_coef) else "-"
        se_display = f"{se:.2f}" if not np.isnan(se) else "-"
        p_display = "< .001" if p_value < 0.001 else f"{p_value:.2f}" if not np.isnan(p_value) else "-"
        sig_display = "**" if p_value < 0.01 else "*" if p_value < 0.05 else "" if not np.isnan(p_value) else ""
        row.cells[1].text = f"{coef_display}{sig_display}"
        row.cells[2].text = std_coef_display
        row.cells[3].text = se_display
        row.cells[4].text = p_display
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].bold = False
            if idx == len(results) - 1:
                set_cell_border(cell, bottom=True)
    
    # Add note with R-squared and adjusted R-squared
    r_squared_display = f"{r_squared:.3f}" if not np.isnan(r_squared) else "-"
    r_squared_adj_display = f"{r_squared_adj:.3f}" if not np.isnan(r_squared_adj) else "-"
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add "Note." (italicized)
    run = note.add_run("Note.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    
    # Add space
    run = note.add_run(" ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Add "R²" (italicized)
    run = note.add_run("R²")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    
    # Add " = {r_squared_display}, "
    run = note.add_run(f" = {r_squared_display}, ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Add "Adjusted" (not italicized)
    run = note.add_run("Adjusted ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Add "R²" (italicized)
    run = note.add_run("R²")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    
    # Add " = {r_squared_adj_display}"
    run = note.add_run(f" = {r_squared_adj_display}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Add model fit note
    f_stat, f_pvalue, df_model, df_resid = diagnostics.get('f_stat', (np.nan, np.nan, np.nan, np.nan))
    f_stat_display = f"{f_stat:.2f}" if not np.isnan(f_stat) else "-"
    f_pvalue_display = "< .001" if f_pvalue < 0.001 else f"{f_pvalue:.3f}" if not np.isnan(f_pvalue) else "-"
    df_model_display = f"{int(df_model)}" if not np.isnan(df_model) else "-"
    df_resid_display = f"{int(df_resid)}" if not np.isnan(df_resid) else "-"
    
    note2 = doc.add_paragraph()
    note2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add "Model: "
    run = note2.add_run("Model: ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Add "F" (italicized)
    run = note2.add_run("F")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    
    # Add "({df_model}, {df_resid}) = {f_stat}, "
    run = note2.add_run(f"({df_model_display}, {df_resid_display}) = {f_stat_display}, ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Add "p" (italicized)
    run = note2.add_run("p")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    
    # Add " = {f_pvalue}"
    run = note2.add_run(f" = {f_pvalue_display}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Add assumptions note
    ad_stat, ad_reject = diagnostics.get('anderson_darling', (np.nan, False))
    ad_stat_display = f"{ad_stat:.2f}" if not np.isnan(ad_stat) else "-"
    ad_reject_display = "Yes" if ad_reject else "No"
    
    bp_lm_stat, bp_pvalue = diagnostics.get('breusch_pagan', (np.nan, np.nan))
    bp_lm_stat_display = f"{bp_lm_stat:.2f}" if not np.isnan(bp_lm_stat) else "-"
    bp_pvalue_display = "< .001" if bp_pvalue < 0.001 else f"{bp_pvalue:.3f}" if not np.isnan(bp_pvalue) else "-"
    
    max_vif = diagnostics.get('max_vif', np.nan)
    max_vif_display = f"{max_vif:.2f}" if not np.isnan(max_vif) else "-"
    
    note3 = doc.add_paragraph()
    note3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add "Assumptions: Anderson-Darling: "
    run = note3.add_run("Assumptions: Anderson-Darling: ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Add "A²" (italicized)
    run = note3.add_run("A²")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    
    # Add " = {ad_stat}, "
    run = note3.add_run(f" = {ad_stat_display}, ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Add "Normality Rejected at 5% = {ad_reject}; "
    run = note3.add_run(f"Normality Rejected at 5% = {ad_reject_display}; ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Add "Breusch-Pagan: "
    run = note3.add_run("Breusch-Pagan: ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Add "LM" (italicized)
    run = note3.add_run("LM")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    
    # Add " = {bp_lm_stat}, "
    run = note3.add_run(f" = {bp_lm_stat_display}, ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Add "p" (italicized)
    run = note3.add_run("p")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    
    # Add " = {bp_pvalue}; "
    run = note3.add_run(f" = {bp_pvalue_display}; ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Add "Max "
    run = note3.add_run("Max ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Add "VIF" (italicized)
    run = note3.add_run("VIF")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    
    # Add " = {max_vif}"
    run = note3.add_run(f" = {max_vif_display}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Add visualizations
    if visualizations.get('qq_plot'):
        doc.add_paragraph("Q-Q Plot of Residuals").alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_picture(BytesIO(base64.b64decode(visualizations['qq_plot'])), width=Inches(4))
    
    if visualizations.get('resid_vs_fitted'):
        doc.add_paragraph("Residuals vs. Fitted Values").alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_picture(BytesIO(base64.b64decode(visualizations['resid_vs_fitted'])), width=Inches(4))
    
    if visualizations.get('resid_histogram'):
        doc.add_paragraph("Histogram of Residuals").alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_picture(BytesIO(base64.b64decode(visualizations['resid_histogram'])), width=Inches(4))
    
    # Save document to a BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Access data from session state
df = st.session_state.get('df', None)
variable_labels = st.session_state.get('variable_labels', {})
value_labels = st.session_state.get('value_labels', {})
visible_columns = st.session_state.get('visible_columns', [])

# Clear session state to avoid caching issues
if 'regression_results' in st.session_state:
    del st.session_state['regression_results']
if 'regression_completed' in st.session_state:
    del st.session_state['regression_completed']

# Initialize session state for this page
if 'regression_results' not in st.session_state:
    st.session_state.regression_results = None
if 'regression_completed' in st.session_state:
    st.session_state.regression_completed = False

# Streamlit UI
st.title("Linear Regression Analysis")

if df is None or df.empty:
    st.warning("No data available. Please upload a dataset on the main page.")
else:
    # Detect all numeric variables (including all PVs for analysis)
    pv_pattern = re.compile(r'^PV([1-9]|10)(MATH|READ|SCIE)(\d*)$', re.IGNORECASE)
    weight_columns = ['W_FSTUWT'] + [f"W_FSTURWT{i}" for i in range(1, 81)] + [col for col in df.columns if 'W_FSCHWT' in col]
    excluded_variables = [
        'CYC', 'NATCEN', 'STRATUM', 'SUBNATIO', 'REGION', 'OECD', 'ADMINMODE',
        'LANGTEST_QQQ', 'LANGTEST_COG', 'LANGTEST_PAQ', 'OPTION_CT', 'OPTION_FL',
        'OPTION_ICTQ', 'OPTION_WBQ', 'OPTION_PQ', 'OPTION_TQ', 'OPTION_UH', 'BOOKID',
        'COBN_S', 'COBN_M', 'COBN_F', 'OCOD1', 'OCOD2', 'OCOD3',
        'ST001D01T', 'ST003D02T', 'ST003D03T',
        'EFFORT1', 'EFFORT2', 'PROGN', 'ISCEDP', 'SENWT', 'VER_DAT', 'TEST',
        'GRADE', 'UNIT', 'WVARSTRR',
        'PAREDINT', 'HISEI', 'HOMEPOS', 'BMMJ1', 'BFMJ2',
        'SCHOOLID', 'STUID', 'CNT', 'CNTRYID', 'CNTSCHID', 'CNTSTUID'
    ]
    item_pattern = re.compile(r'^(ST|FL|IC|WB|PA)\w{8}$', re.IGNORECASE)
    numeric_vars = [
        col for col in df.columns 
        if col not in weight_columns 
        and col not in excluded_variables
        and not item_pattern.match(col) 
        and not df[col].isna().all()
        and df[col].dtype in ['float64', 'int64']
    ]
    
    # Identify plausible value domains and regular numeric variables
    pv_domains = {}
    regular_numeric_vars = []
    for col in numeric_vars:
        pv_match = pv_pattern.match(col)
        if pv_match:
            pv_num = int(pv_match.group(1))
            domain = pv_match.group(2).upper()  # Normalize to uppercase
            if domain not in pv_domains:
                pv_domains[domain] = []
            pv_domains[domain].append(col)
        else:
            regular_numeric_vars.append(col)
    
    # Sort PVs within each domain to ensure PV1 to PV10 order
    for domain in pv_domains:
        pv_domains[domain].sort(key=lambda x: int(re.match(r'PV(\d+)', x, re.IGNORECASE).group(1)))
    
    # Check if PV domains have all 10 plausible values
    for domain, pv_list in pv_domains.items():
        if len(pv_list) != 10:
            st.warning(f"Domain {domain} has {len(pv_list)} plausible values instead of 10. Only {pv_list} will be used.")
    
    # Create domain options for selection
    domain_to_label = {
        'MATH': 'Mathematics score',
        'READ': 'Reading score',
        'SCIE': 'Science score'
    }
    domain_options = [domain_to_label.get(domain, domain) for domain in pv_domains.keys()]
    label_to_domain = {domain_to_label.get(domain, domain): domain for domain in pv_domains.keys()}
    
    # Prepare labels for regular numeric variables, using only visible ones for UI
    regular_numeric_vars = [col for col in regular_numeric_vars if col in visible_columns]
    var_labels = [variable_labels.get(col, col) for col in regular_numeric_vars]
    label_to_var = {variable_labels.get(col, col): col for col in regular_numeric_vars}
    unique_var_labels = []
    seen_labels = {}
    for col, label in zip(regular_numeric_vars, var_labels):
        if label in seen_labels:
            unique_label = f"{label} ({col})"
        else:
            unique_label = label
        seen_labels[label] = True
        unique_var_labels.append(unique_label)
        label_to_var[unique_label] = col
    
    # Combine domain options and regular numeric variables for the dropdown
    all_var_options = domain_options + unique_var_labels
    
    if len(all_var_options) < 2:
        st.warning("At least two numeric variables or domains are required for linear regression analysis.")
    else:
        # Select Dependent Variable
        st.write("Select the dependent variable:")
        dependent_var_label = st.selectbox(
            "Dependent Variable",
            all_var_options,
            index=0 if "Mathematics score" not in all_var_options else all_var_options.index("Mathematics score"),
            key="dep_var"
        )
        
        # Determine if the dependent variable is a domain or a regular variable
        if dependent_var_label in label_to_domain:
            dependent_var = pv_domains[label_to_domain[dependent_var_label]][0]  # Use the first PV (e.g., PV1MATH)
        else:
            dependent_var = label_to_var[dependent_var_label]
        
        # Select Covariates (default to ST004D01T and ESCS)
        st.write("Select covariates (default: Student gender, ESCS):")
        covariate_options = [label for label in all_var_options if label != dependent_var_label]
        # Map ST004D01T and ESCS to their labels
        default_covariates = []
        student_gender_label = variable_labels.get("ST004D01T", "ST004D01T")
        # Check if the label for ST004D01T was modified due to duplicates
        for label in unique_var_labels:
            if label == student_gender_label or label.startswith(f"{student_gender_label} ("):
                student_gender_label = label
                break
        escs_label = variable_labels.get("ESCS", "ESCS")
        for label in unique_var_labels:
            if label == escs_label or label.startswith(f"{escs_label} ("):
                escs_label = label
                break
        # Set default covariates
        if student_gender_label in covariate_options:
            default_covariates.append(student_gender_label)
        if escs_label in covariate_options:
            default_covariates.append(escs_label)
        
        covariate_labels = st.multiselect(
            "Covariates",
            covariate_options,
            default=default_covariates,
            key="covariates"
        )
        
        # Map covariate labels to actual column names
        covariates = []
        for label in covariate_labels:
            if label in label_to_domain:
                # Use the first PV for the domain (e.g., PV1MATH for Mathematics score)
                domain = label_to_domain[label]
                covariates.append(pv_domains[domain][0])
            else:
                covariates.append(label_to_var[label])
        
        # Select Predictors (optional, blank by default)
        st.write("Select predictors (optional):")
        predictor_options = [label for label in all_var_options if label != dependent_var_label and label not in covariate_labels]
        predictor_labels = st.multiselect(
            "Predictors (Optional)",
            predictor_options,
            default=[],
            key="predictors"
        )
        
        # Map predictor labels to actual column names
        predictors = []
        for label in predictor_labels:
            if label in label_to_domain:
                domain = label_to_domain[label]
                predictors.append(pv_domains[domain][0])
            else:
                predictors.append(label_to_var[label])
        
        # Combine covariates and predictors for the regression
        independent_vars = covariates + predictors
        
        run_analysis = st.button("Run Analysis", key="run_regression")
        
        if run_analysis and independent_vars:
            try:
                if 'W_FSTUWT' not in df.columns:
                    st.error("Final student weight (W_FSTUWT) not found in the dataset.")
                else:
                    # Check for replicate weights availability
                    replicate_weight_cols = [f"W_FSTURWT{i}" for i in range(1, 81)]
                    missing_weights = [col for col in replicate_weight_cols if col not in df.columns]
                    use_brr = len(missing_weights) == 0
                    
                    # Create a placeholder for status messages
                    status_placeholder = st.empty()
                    
                    # Compute regression with PVs and BRR, passing label_to_var for mapping and the placeholder
                    results, r_squared, r_squared_adj, diagnostics, visualizations = compute_linear_regression_with_pvs(
                        df, dependent_var, independent_vars, df['W_FSTUWT'], replicate_weight_cols, use_brr, label_to_var, status_placeholder
                    )
                    
                    # Store results in session state
                    st.session_state.regression_results = results
                    st.session_state.regression_completed = True
                    st.session_state.r_squared = r_squared
                    st.session_state.r_squared_adj = r_squared_adj
                    st.session_state.diagnostics = diagnostics
                    st.session_state.visualizations = visualizations
                    
                    # Render the regression table
                    status_placeholder.write("Rendering regression table...")
                    table_html = render_regression_table(dependent_var_label, results, r_squared, r_squared_adj, diagnostics)
                    components.html(table_html, height=400, scrolling=True)
                    
                    # Provide download button for Word document
                    doc_buffer = create_word_table(dependent_var_label, results, r_squared, r_squared_adj, diagnostics, visualizations)
                    st.download_button(
                        label="Download Table as Word Document",
                        data=doc_buffer,
                        file_name="Linear_Regression_Table.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    # Display visualizations without headers
                    if visualizations.get('qq_plot'):
                        st.image(f"data:image/png;base64,{visualizations['qq_plot']}")
                    
                    if visualizations.get('resid_vs_fitted'):
                        st.image(f"data:image/png;base64,{visualizations['resid_vs_fitted']}")
                    
                    if visualizations.get('resid_histogram'):
                        st.image(f"data:image/png;base64,{visualizations['resid_histogram']}")
                    
                    status_placeholder.write("Linear regression analysis completed.")
                    status_placeholder.empty()  # Clear the placeholder after completion
            except Exception as e:
                st.error(f"Error computing linear regression: {str(e)}")
                st.session_state.regression_completed = False
                st.session_state.regression_results = None
        elif st.session_state.regression_results and st.session_state.regression_completed:
            if independent_vars:
                results = st.session_state.regression_results
                r_squared = st.session_state.get('r_squared', np.nan)
                r_squared_adj = st.session_state.get('r_squared_adj', np.nan)
                diagnostics = st.session_state.get('diagnostics', {})
                visualizations = st.session_state.get('visualizations', {})
                table_html = render_regression_table(dependent_var_label, results, r_squared, r_squared_adj, diagnostics)
                components.html(table_html, height=400, scrolling=True)
                
                # Provide download button for Word document
                doc_buffer = create_word_table(dependent_var_label, results, r_squared, r_squared_adj, diagnostics, visualizations)
                st.download_button(
                    label="Download Table as Word Document",
                    data=doc_buffer,
                    file_name="Linear_Regression_Table.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                # Display visualizations without headers
                if visualizations.get('qq_plot'):
                    st.image(f"data:image/png;base64,{visualizations['qq_plot']}")
                
                if visualizations.get('resid_vs_fitted'):
                    st.image(f"data:image/png;base64,{visualizations['resid_vs_fitted']}")
                
                if visualizations.get('resid_histogram'):
                    st.image(f"data:image/png;base64,{visualizations['resid_histogram']}")
                
                st.write("Linear regression analysis completed.")
            else:
                st.write("Please select at least one covariate to compute linear regression.")
        else:
            st.write("Please select a dependent variable and at least one covariate, then click 'Run Analysis' to compute linear regression.")

# Instructions section
st.header("Instructions")
st.markdown("""
- **Select Variables**: Choose a dependent variable, covariates (default: Student gender, ESCS), and optional predictors from the dropdown menus. Plausible value domains (e.g., Mathematics score, Reading score) will use all 10 plausible values for analysis, combined using Rubin's rules.
- **Run Analysis**: Click "Run Analysis" to perform the weighted linear regression analysis using student weights (W_FSTUWT). Standard errors will be computed using replicate weights if available.
- **View Results**: Results are displayed in an APA-style table with coefficients, standard errors, and p-values rounded to 2 decimal places. You can download the table as a Word document using the download button.
- **Navigate**: Use the sidebar to switch between different analysis types or return to the main page to upload a new dataset.
""")