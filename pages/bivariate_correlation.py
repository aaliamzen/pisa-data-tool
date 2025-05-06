import streamlit as st
import pandas as pd
import numpy as np
import re
import streamlit.components.v1 as components
from scipy.stats import t
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Streamlit app configuration
st.set_page_config(page_title="Bivariate Correlation - PISA Data Exploration Tool", layout="wide")

# Function to compute weighted correlation coefficient
def weighted_correlation(x, y, w):
    try:
        if len(x) != len(y) or len(x) != len(w):
            raise ValueError("Input arrays must have the same length.")
        if not all(w >= 0):
            raise ValueError("Weights must be non-negative.")
        if len(x) <= 1:
            return np.nan, np.nan
        # Remove NaN values
        mask = ~(np.isnan(x) | np.isnan(y) | np.isnan(w))
        x = x[mask]
        y = y[mask]
        w = w[mask]
        if len(x) <= 1:
            return np.nan, np.nan
        # Weighted mean
        w_sum = np.sum(w)
        if w_sum == 0:
            return np.nan, np.nan
        wx_mean = np.sum(w * x) / w_sum
        wy_mean = np.sum(w * y) / w_sum
        # Weighted covariance and variances
        cov_xy = np.sum(w * (x - wx_mean) * (y - wy_mean)) / w_sum
        var_x = np.sum(w * (x - wx_mean)**2) / w_sum
        var_y = np.sum(w * (y - wy_mean)**2) / w_sum
        if var_x == 0 or var_y == 0:
            return np.nan, np.nan
        # Weighted correlation
        corr = cov_xy / np.sqrt(var_x * var_y)
        # Compute p-value (approximate, using effective sample size)
        n_eff = len(x)  # Simplified; could use more sophisticated effective sample size
        if n_eff < 3:
            return corr, np.nan
        t_stat = corr * np.sqrt((n_eff - 2) / (1 - corr**2))
        p_value = 2 * (1 - t.cdf(np.abs(t_stat), df=n_eff-2))
        return corr, p_value
    except Exception as e:
        st.error(f"Error in weighted_correlation: {str(e)}")
        return np.nan, np.nan

# Function to compute BRR standard errors for correlation
def compute_brr_se_correlation(x, y, replicate_weights, corr_data, progress_bar):
    try:
        # Initial correlation with final student weights for reference
        main_corr, _ = weighted_correlation(x, y, corr_data['W_FSTUWT'].values)
        if np.isnan(main_corr):
            return np.nan
        
        # Compute correlations for each replicate weight
        replicate_corrs = []
        total_weights = len(replicate_weights)
        for idx, weight_col in enumerate(replicate_weights):
            # Create a DataFrame with only the weights to handle NA dropping
            data = corr_data[[weight_col, 'W_FSTUWT']].copy()
            data = data.dropna()
            if len(data) < 2:
                continue
            
            # Subset x and y to match the non-NA indices of the weights
            indices = data.index
            pos = corr_data.index.get_indexer(indices)
            pos = pos[pos != -1]  # Remove invalid indices
            if len(pos) < 2:
                continue
            x_rep = x[pos]
            y_rep = y[pos]
            w_rep = data[weight_col].values
            
            # Compute correlation with replicate weights
            rep_corr, _ = weighted_correlation(x_rep, y_rep, w_rep)
            if not np.isnan(rep_corr):
                replicate_corrs.append(rep_corr)
            # Update progress bar
            progress_bar.progress((idx + 1) / total_weights)
        
        if not replicate_corrs:
            return np.nan
        
        # Convert to array and compute standard error
        replicate_corrs = np.array(replicate_corrs)
        se = np.sqrt((1 / 80) * np.sum((replicate_corrs - main_corr) ** 2))
        return se
    except Exception as e:
        st.error(f"Error in compute_brr_se_correlation: {str(e)}")
        return np.nan

# Function to compute p-values using BRR standard errors
def compute_brr_p_value(corr, se, n):
    try:
        if np.isnan(se) or se == 0:
            return np.nan
        t_stat = corr / se
        p_value = 2 * (1 - t.cdf(np.abs(t_stat), df=n-2))
        return p_value
    except Exception as e:
        st.error(f"Error in compute_brr_p_value: {str(e)}")
        return np.nan

# Function to apply Rubin's rules for combining correlations across plausible values
def apply_rubins_rules_correlations(corrs_list, se_list, n):
    try:
        corrs_array = np.array(corrs_list)
        se_array = np.array(se_list)
        num_pvs = len(corrs_list)
        
        # Combined point estimate
        combined_corr = np.mean(corrs_array)
        
        # Within-imputation variance
        within_var = np.mean(se_array**2)
        
        # Between-imputation variance
        between_var = np.var(corrs_array, ddof=1)
        
        # Total variance
        total_var = within_var + (1 + 1/num_pvs) * between_var
        
        # Combined standard error
        combined_se = np.sqrt(total_var)
        
        # Compute p-value
        p_value = compute_brr_p_value(combined_corr, combined_se, n)
        
        return combined_corr, combined_se, p_value
    except Exception as e:
        st.error(f"Error in apply_rubins_rules_correlations: {str(e)}")
        return np.nan, np.nan, np.nan

# Function to render correlation table as HTML
def render_correlation_table(selected_labels, selected_codes, results):
    html_content = """
    <style>
    .corr-table-container {
        display: inline-block;
        overflow-x: auto;
        scrollbar-width: thin;
        min-width: 0;
        margin: 20px 0;
    }
    .corr-table-container::-webkit-scrollbar {
        height: 8px;
    }
    .corr-table-container::-webkit-scrollbar-thumb {
        background-color: #888;
        border-radius: 4px;
    }
    .corr-table {
        table-layout: fixed;
        border-collapse: collapse;
        font-size: 14px;
        margin: 0;
    }
    .corr-table th, .corr-table td {
        border: none;
        padding: 8px;
        box-sizing: border-box;
        text-align: center;
    }
    .corr-table th:first-child, .corr-table td:first-child {
        width: 200px !important;
        text-align: left;
        white-space: normal;
        overflow-wrap: break-word;
    }
    .corr-table th:not(:first-child), .corr-table td:not(:first-child) {
        width: 100px !important;
    }
    .corr-table tr:nth-child(even) {
        background-color: #f9f9f9;
    }
    .corr-table-title {
        font-size: 16px;
        font-weight: bold;
        text-align: left;
        margin-bottom: 5px;
    }
    .corr-table-subtitle {
        font-size: 16px;
        font-style: italic;
        text-align: left;
        margin-bottom: 10px;
    }
    .corr-table-header {
        border-top: 1px solid #000;
        border-bottom: 1px solid #000;
    }
    .corr-table-last-row {
        border-bottom: 1px solid #000;
    }
    </style>
    <div class="corr-table-container">
        <div class="corr-table-title">Table 1</div>
        <div class="corr-table-subtitle">Weighted Bivariate Correlations Between Selected Variables</div>
        <table class="corr-table">
            <tr class="corr-table-header">
                <th>Variable 1</th>
                <th>Variable 2</th>
                <th>Correlation</th>
                <th>P-value</th>
                <th>Standard Error</th>
            </tr>
            {{data_rows}}
        </table>
    </div>
    """
    data_rows = ""
    for idx, (var1_label, var2_label, corr, p_value, se) in enumerate(results):
        corr_display = f"{corr:.3f}" if not np.isnan(corr) else "-"
        p_display = "< .001" if p_value < 0.001 else f"= {p_value:.3f}" if not np.isnan(p_value) else "-"
        se_display = f"{se:.3f}" if not np.isnan(se) else "-"
        sig_display = "**" if p_value < 0.01 else "*" if p_value < 0.05 else "" if not np.isnan(p_value) else ""
        row_class = "corr-table-last-row" if idx == len(results) - 1 else ""
        row = f"""
        <tr class="{row_class}">
            <th>{var1_label}</th>
            <th>{var2_label}</th>
            <td>{corr_display}{sig_display}</td>
            <td>{p_display}</td>
            <td>{se_display}</td>
        </tr>
        """
        data_rows += row
    
    full_html = html_content.replace("{{data_rows}}", data_rows)
    return full_html

# Function to set cell borders in Word document
def set_cell_border(cell, top=False, bottom=False, left=False, right=False):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    # Remove existing borders
    for border in ['top', 'bottom', 'left', 'right']:
        existing_border = tcPr.find(qn(f'w:{border}'))
        if existing_border is not None:
            tcPr.remove(existing_border)
    
    # Set new borders where specified
    if top:
        border = OxmlElement('w:top')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Border width in eighths of a point
        tcPr.append(border)
    if bottom:
        border = OxmlElement('w:bottom')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        tcPr.append(border)
    if left:
        border = OxmlElement('w:left')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        tcPr.append(border)
    if right:
        border = OxmlElement('w:right')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        tcPr.append(border)

# Function to create a Word document with the APA-style table
def create_word_table(results):
    doc = Document()
    
    # Set document to landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    
    # Add title
    title = doc.add_paragraph("Table 1")
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.runs[0].font.name = 'Times New Roman'
    title.runs[0].font.size = Pt(10)
    
    # Add subtitle
    subtitle = doc.add_paragraph("Weighted Bivariate Correlations Between Selected Variables")
    subtitle.runs[0].italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle.runs[0].font.name = 'Times New Roman'
    subtitle.runs[0].font.size = Pt(10)
    
    # Create table
    table = doc.add_table(rows=1 + len(results), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Normal Table'  # Use a style with no borders
    
    # Set column widths
    for column in table.columns:
        for cell in column.cells:
            cell.width = Inches(1.5)  # Adjusted for landscape orientation
    
    # Add header row
    headers = ["Variable 1", "Variable 2", "Correlation", "P-value", "Standard Error"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].bold = True
        set_cell_border(cell, top=True, bottom=True)
    
    # Add data rows
    for idx, (var1_label, var2_label, corr, p_value, se) in enumerate(results):
        row = table.rows[idx + 1]
        row.cells[0].text = var1_label
        row.cells[1].text = var2_label
        corr_display = f"{corr:.3f}" if not np.isnan(corr) else "-"
        p_display = "< .001" if p_value < 0.001 else f"= {p_value:.3f}" if not np.isnan(p_value) else "-"
        se_display = f"{se:.3f}" if not np.isnan(se) else "-"
        sig_display = "**" if p_value < 0.01 else "*" if p_value < 0.05 else "" if not np.isnan(p_value) else ""
        row.cells[2].text = f"{corr_display}{sig_display}"
        row.cells[3].text = p_display
        row.cells[4].text = se_display
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if idx == len(results) - 1:
                set_cell_border(cell, bottom=True)
    
    # Save document to a BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Access data from session state
df = st.session_state.get('df', None)
variable_labels = st.session_state.get('variable_labels', {})
value_labels = st.session_state.get('value_labels', {})
visible_columns = st.session_state.get('visible_columns', [])

# Initialize session state for this page
if 'bivariate_results' not in st.session_state:
    st.session_state.bivariate_results = None
if 'bivariate_completed' not in st.session_state:
    st.session_state.bivariate_completed = False
if 'bivariate_selected_codes' not in st.session_state:
    st.session_state.bivariate_selected_codes = None

# Streamlit UI
st.title("Bivariate Correlation Analysis")

if df is None or df.empty:
    st.warning("No data available. Please upload a dataset on the main page.")
else:
    # Detect all numeric variables (including all PVs for analysis)
    pv_pattern = re.compile(r'^PV([1-9]|10)(MATH|READ|SCIE)(\d*)$', re.IGNORECASE)
    weight_columns = ['W_FSTUWT'] + [f"W_FSTURWT{i}" for i in range(1, 81)] + [col for col in df.columns if 'W_FSCHWT' in col]
    excluded_variables = [
        'CYC', 'NATCEN', 'STRATUM', 'SUBNATIO', 'REGION', 'OECD', 'ADMINMODE',
        'LANGTEST_QQQ', 'LANGTEST_COG', 'LANGTEST_PAQ', 'OPTION_CT', 'OPTION_FL',
        'OPTION_ICTQ', 'OPTION_WBQ', 'OPTION_PQ', 'OPTION_TQ', 'OPTION_UH', 'BOOKID',
        'COBN_S', 'COBN_M', 'COBN_F', 'OCOD1', 'OCOD2', 'OCOD3',
        'ST001D01T', 'ST003D02T', 'ST003D03T',
        'EFFORT1', 'EFFORT2', 'PROGN', 'ISCEDP', 'SENWT', 'VER_DAT', 'TEST',
        'GRADE', 'UNIT', 'WVARSTRR',
        'PAREDINT', 'HISEI', 'HOMEPOS', 'BMMJ1', 'BFMJ2',
        'SCHOOLID', 'STUID', 'CNT', 'CNTRYID', 'CNTSCHID', 'CNTSTUID'
    ]
    item_pattern = re.compile(r'^(ST|FL|IC|WB|PA)\w{8}$', re.IGNORECASE)
    numeric_vars = [
        col for col in df.columns 
        if col not in weight_columns 
        and col not in excluded_variables
        and not item_pattern.match(col) 
        and not df[col].isna().all()
        and df[col].dtype in ['float64', 'int64']
    ]
    
    # Identify plausible value domains and regular numeric variables
    pv_domains = {}
    regular_numeric_vars = []
    for col in numeric_vars:
        pv_match = pv_pattern.match(col)
        if pv_match:
            pv_num = int(pv_match.group(1))
            domain = pv_match.group(2).upper()  # Normalize to uppercase
            if domain not in pv_domains:
                pv_domains[domain] = []
            pv_domains[domain].append(col)
        else:
            regular_numeric_vars.append(col)
    
    # Sort PVs within each domain to ensure PV1 to PV10 order
    for domain in pv_domains:
        pv_domains[domain].sort(key=lambda x: int(re.match(r'PV(\d+)', x, re.IGNORECASE).group(1)))
    
    # Check if PV domains have all 10 plausible values
    for domain, pv_list in pv_domains.items():
        if len(pv_list) != 10:
            st.warning(f"Domain {domain} has {len(pv_list)} plausible values instead of 10. Only {pv_list} will be used.")
    
    # Create domain options for selection
    domain_to_label = {
        'MATH': 'Mathematics score',
        'READ': 'Reading score',
        'SCIE': 'Science score'
    }
    domain_options = [domain_to_label.get(domain, domain) for domain in pv_domains.keys()]
    label_to_domain = {domain_to_label.get(domain, domain): domain for domain in pv_domains.keys()}
    
    # Prepare labels for regular numeric variables, using only visible ones for UI
    regular_numeric_vars = [col for col in regular_numeric_vars if col in visible_columns]
    var_labels = [variable_labels.get(col, col) for col in regular_numeric_vars]
    label_to_var = {variable_labels.get(col, col): col for col in regular_numeric_vars}
    unique_var_labels = []
    seen_labels = {}
    for col, label in zip(regular_numeric_vars, var_labels):
        if label in seen_labels:
            unique_label = f"{label} ({col})"
        else:
            unique_label = label
        seen_labels[label] = True
        unique_var_labels.append(unique_label)
        label_to_var[unique_label] = col
    
    if len(unique_var_labels) + len(domain_options) < 2:
        st.warning("At least two numeric variables or domains are required for bivariate correlation analysis.")
    else:
        # Select Variables (default to MATH domain and ESCS)
        st.write("Select variables for bivariate correlation (default: Mathematics score, ESCS):")
        all_var_options = domain_options + unique_var_labels
        default_vars = []
        if "Mathematics score" in all_var_options:
            default_vars.append("Mathematics score")
        escs_label = variable_labels.get("ESCS", "ESCS")
        if escs_label in seen_labels and escs_label in all_var_options:
            default_vars.append(escs_label)
        elif all_var_options and len(all_var_options) > 1:
            default_vars.append(all_var_options[1 if "Mathematics score" in all_var_options else 0])
        
        if len(default_vars) < 2 and len(all_var_options) >= 2:
            default_vars = all_var_options[:2]
        
        selected_var_labels = st.multiselect(
            "Variables",
            all_var_options,
            default=default_vars,
            key="bivariate_vars"
        )
        
        # Split selected variables into domains and regular variables
        selected_domains = []
        selected_vars = []
        selected_codes = []
        all_vars = []
        for label in selected_var_labels:
            if label in label_to_domain:
                domain = label_to_domain[label]
                selected_domains.append(domain)
                selected_codes.append(domain)
                all_vars.append(pv_domains[domain])
            else:
                selected_vars.append(label_to_var[label])
                selected_codes.append(label_to_var[label])
                all_vars.append([label_to_var[label]])
        
        run_analysis = st.button("Run Analysis", key="run_bivariate")
        
        if run_analysis and (selected_vars or selected_domains) and len(selected_var_labels) >= 2:
            try:
                if 'W_FSTUWT' not in df.columns:
                    st.error("Final student weight (W_FSTUWT) not found in the dataset.")
                else:
                    # Check for replicate weights availability
                    replicate_weight_cols = [f"W_FSTURWT{i}" for i in range(1, 81)]
                    missing_weights = [col for col in replicate_weight_cols if col not in df.columns]
                    use_brr = len(missing_weights) == 0
                    
                    # Compute the maximum number of plausible values
                    max_pvs = max([len(vars) for vars in all_vars])
                    
                    # Progress bar for PV iterations
                    pv_progress = st.progress(0)
                    total_iterations = (len(all_vars) * (len(all_vars) - 1)) // 2 * max_pvs
                    iteration_count = 0
                    
                    # Compute correlations for each pair of variables
                    results = []
                    for i in range(len(all_vars)):
                        for j in range(i + 1, len(all_vars)):
                            var1_list = all_vars[i]
                            var2_list = all_vars[j]
                            var1_is_pv = len(var1_list) > 1
                            var2_is_pv = len(var2_list) > 1
                            num_pvs = max(len(var1_list), len(var2_list))
                            
                            all_corrs = []
                            all_p_values = []
                            all_se = []
                            
                            for pv_idx in range(num_pvs):
                                var1 = var1_list[pv_idx % len(var1_list)]
                                var2 = var2_list[pv_idx % len(var2_list)]
                                st.write(f"Processing correlation between {var1} and {var2}...")
                                
                                # Create a subset DataFrame for just this pair
                                pair_columns = ['W_FSTUWT']
                                if use_brr:
                                    pair_columns += replicate_weight_cols
                                pair_columns.append(var1)
                                pair_columns.append(var2)
                                pair_data = df[pair_columns].dropna()
                                
                                if len(pair_data) < 2:
                                    st.warning(f"Insufficient non-missing data for correlation between {var1} and {var2}.")
                                    all_corrs.append(np.nan)
                                    all_p_values.append(np.nan)
                                    all_se.append(np.nan)
                                    iteration_count += 1
                                    pv_progress.progress(iteration_count / total_iterations)
                                    continue
                                
                                x = pair_data[var1].values
                                y = pair_data[var2].values
                                w = pair_data['W_FSTUWT'].values
                                corr, p_value = weighted_correlation(x, y, w)
                                if use_brr:
                                    st.write("Calculating standard error using replicate weights...")
                                    brr_progress = st.progress(0)
                                    se = compute_brr_se_correlation(x, y, replicate_weight_cols, pair_data, brr_progress)
                                    p_value = compute_brr_p_value(corr, se, len(pair_data))
                                else:
                                    se = np.nan
                                
                                all_corrs.append(corr)
                                all_p_values.append(p_value)
                                all_se.append(se)
                                
                                iteration_count += 1
                                pv_progress.progress(iteration_count / total_iterations)
                        
                            # Combine results
                            if num_pvs == 1:
                                st.write(f"Combining results for {selected_var_labels[i]} and {selected_var_labels[j]} (single iteration, no plausible values)...")
                                combined_corr = all_corrs[0]
                                combined_p_value = all_p_values[0]
                                combined_se = all_se[0]
                            else:
                                st.write(f"Combining results for {selected_var_labels[i]} and {selected_var_labels[j]}...")
                                combined_corr, combined_se, combined_p_value = apply_rubins_rules_correlations(all_corrs, all_se, len(pair_data))
                            
                            # Store results
                            results.append((selected_var_labels[i], selected_var_labels[j], combined_corr, combined_p_value, combined_se))
                    
                    # Store results and selected_codes in session state
                    st.session_state.bivariate_results = results
                    st.session_state.bivariate_selected_codes = selected_codes
                    st.session_state.bivariate_completed = True
                    
                    # Render the correlation table
                    st.write("Rendering correlation table...")
                    table_html = render_correlation_table(selected_var_labels, selected_codes, results)
                    components.html(table_html, height=400, scrolling=True)
                    st.write("Bivariate correlation analysis completed.")
                    
                    # Provide download button for Word document
                    doc_buffer = create_word_table(results)
                    st.download_button(
                        label="Download Table as Word Document",
                        data=doc_buffer,
                        file_name="Bivariate_Correlation_Table.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            except Exception as e:
                st.error(f"Error computing bivariate correlations: {str(e)}")
                st.session_state.bivariate_completed = False
                st.session_state.bivariate_results = None
                st.session_state.bivariate_selected_codes = None
        elif st.session_state.bivariate_results and st.session_state.bivariate_completed:
            if selected_var_labels:
                results = st.session_state.bivariate_results
                selected_codes = st.session_state.bivariate_selected_codes
                table_html = render_correlation_table(selected_var_labels, selected_codes, results)
                components.html(table_html, height=400, scrolling=True)
                st.write("Bivariate correlation analysis completed.")
                
                # Provide download button for Word document
                doc_buffer = create_word_table(results)
                st.download_button(
                    label="Download Table as Word Document",
                    data=doc_buffer,
                    file_name="Bivariate_Correlation_Table.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.write("Please select at least two variables to compute bivariate correlations.")
        else:
            st.write("Please select at least two variables and click 'Run Analysis' to compute bivariate correlations.")

# Instructions section
st.header("Instructions")
st.markdown("""
- **Select Variables**: Choose two or more variables (domains or numeric variables) from the dropdown menus. Plausible value domains (e.g., Mathematics score, Reading score) will use all 10 plausible values for analysis.
- **Run Analysis**: Click "Run Analysis" to perform the weighted bivariate correlation analysis. Analyses involving plausible values will be combined using Rubin's rules.
- **View Results**: Results are displayed in an APA-style table. You can download the table as a Word document using the download button.
- **Navigate**: Use the sidebar to switch between different analysis types or return to the main page to upload a new dataset.
""")